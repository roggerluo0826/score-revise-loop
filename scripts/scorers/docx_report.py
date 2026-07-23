# -*- coding: utf-8 -*-
"""打分器:節能報告 / ESCO 計畫書 docx。

檢查項全部來自實際踩過的坑(報告拷貝忍術第 7 步防呆清單 + 歷次案件)。
每一項都必須是「程式能判定對錯」的,主觀的一律不放進來。
"""
from __future__ import annotations

import re
import zipfile

import docx
from docx.oxml.ns import qn

from scorelib import Report, parse_num, close

# 母版拷貝後最容易殘留的機構名。先抓字尾關鍵字,再往前吃字直到遇到 STOP,
# 直接用貪婪 regex 會把「汰換高中」「以及國民中學」這種動詞/連接詞前綴一起吃進來。
ORG_SUFFIX = ('股份有限公司', '有限公司', '高級工業職業學校', '工業職業學校',
              '高級中學', '國民中學', '國民小學', '護理之家', '養護中心',
              '科技大學', '高工', '高中', '醫院', '診所')
ORG_SUFFIX_RE = re.compile('|'.join(ORG_SUFFIX))
# 不可能出現在機構名內部的字(助詞、連接詞、動詞、空間量詞),往前吃字時遇到就停。
# 注意:不要放「高、新、中、大」這種會出現在校名裡的字,否則「高雄市立…」會被截斷。
STOP_CHARS = set('的了和與及或並在為是有本該此其之等就都也將把從對向再又最很如若則會'
                 '可要能不無需應以由於受經按依請詳見另計約至含共總合案例各所但而故因'
                 '我你他們貴汰換修復棟樓層間室處組課科第台部區段項台座式型')
# 泛稱,不足以判定是某一家機構
GENERIC = set(ORG_SUFFIX)
# 常黏在機構名前面的修飾語,抓到後從左邊剝掉
LEAD_NOISE = ('輔導團隊', '執行單位', '委託單位', '受託單位', '主辦單位', '本案', '該案')
# 預設放行:輔導團隊自己的名字本來就會出現在報告裡
DEFAULT_ALLOW = ()          # 用 --allow 帶入執行單位自己的公司名
MIN_ORG_LEN = 4


def extract_orgs(text):
    """回傳 {機構名: 次數}。先抓字尾關鍵字再往前吃字,比貪婪 regex 準。"""
    found = {}
    for m in ORG_SUFFIX_RE.finditer(text):
        i, j = m.start(), m.start()
        while j > 0 and (m.end() - (j - 1)) <= 14:
            ch = text[j - 1]
            if not ('一' <= ch <= '鿿') or ch in STOP_CHARS:
                break
            j -= 1
        name = text[j:m.end()]
        changed = True
        while changed:                       # 反覆剝掉左邊的修飾語
            changed = False
            for noise in LEAD_NOISE:
                if name.startswith(noise) and len(name) > len(noise):
                    name, changed = name[len(noise):], True
        if name in GENERIC or len(name) < MIN_ORG_LEN:
            continue
        found[name] = found.get(name, 0) + 1
    return found
PLACEHOLDER_RE = re.compile(r'待補|待確認|TBD|ＴＢＤ|XXX|ＸＸＸ|\?\?\?|？？？|<[^>]{1,20}>')
NUM_RE = re.compile(r'\d{1,3}(?:,\d{3})+(?:\.\d+)?|\d+(?:\.\d+)?')


# ---------- 表格定位:用內容特徵找,不用固定索引 ----------
def find_tables(d):
    """回傳 {summary, basic, equip, measures[]}"""
    out = dict(summary=None, basic=None, equip=None, measures=[])
    for t in d.tables:
        head = ' '.join(c.text for r in t.rows[:8] for c in r.cells)
        ncol = len(t.columns)
        if out['summary'] is None and ncol == 6 and '項次' in head and '建議改善措施名稱' in head:
            out['summary'] = t
        elif out['basic'] is None and '單位名稱' in head and ('契約' in head or '台電' in head):
            out['basic'] = t
        elif out['equip'] is None and '財產名稱' in head:
            out['equip'] = t
        elif '系統別' in head and '改善措施' in head:
            out['measures'].append(t)
    return out


def cells_of(row):
    """去掉合併儲存格造成的重複,回傳 (欄索引, 文字) 清單。"""
    seen, out = [], []
    for i, c in enumerate(row.cells):
        if any(p is c._tc for p in seen):
            continue
        seen.append(c._tc)
        out.append((i, c.text.strip()))
    return out


def all_paragraphs(d):
    def walk_cell(cell):
        for p in cell.paragraphs:
            yield p
        for t in cell.tables:
            yield from walk_table(t)

    def walk_table(t):
        seen = []
        for row in t.rows:
            for c in row.cells:
                if any(p is c._tc for p in seen):
                    continue
                seen.append(c._tc)
                yield from walk_cell(c)

    for p in d.paragraphs:
        yield p
    for t in d.tables:
        yield from walk_table(t)
    for sec in d.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                yield p


def run_fmt(r):
    rPr = r._element.find(qn('w:rPr'))
    asc = ea = None
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            asc, ea = rf.get(qn('w:ascii')), rf.get(qn('w:eastAsia'))
    sz = r.font.size.pt if r.font.size else None
    return (asc, ea, sz)


# =========================== 主檢查 ===========================
def score(path, forbid=(), own_name=None, min_chars=200, allow=(), **_):
    rep = Report(target=path, scorer='docx_report')
    d = docx.Document(path)
    T = find_tables(d)
    rep.meta['表格'] = (f"摘要表{'有' if T['summary'] is not None else '無'} / "
                      f"基本資料{'有' if T['basic'] is not None else '無'} / "
                      f"設備表{'有' if T['equip'] is not None else '無'} / "
                      f"改善措施表 {len(T['measures'])} 張")

    _check_stale(d, rep, forbid, own_name, allow)
    _check_placeholder(d, rep)
    _check_halfwidth_comma(d, rep)
    _check_font_uniformity(d, rep)
    _check_heading_style(d, rep)
    _check_toc(path, rep)
    ms = _check_measures(T, rep, min_chars)
    _check_summary(T, rep, ms)
    return rep


def _check_stale(d, rep, forbid, own_name, allow=()):
    rep.ran('stale_client')
    text = '\n'.join(p.text for p in all_paragraphs(d))
    for f in forbid:
        f = f.strip()
        if f and f in text:
            i = text.find(f)
            rep.add('stale_client', 'blocker', f'殘留指定的舊客戶字串「{f}」',
                    where=repr(text[max(0, i - 20):i + 25]),
                    fix=f'全文搜尋「{f}」並改成本案對應內容')
    # 自動偵測只是提示,不是判決:regex 認機構名一定有誤差,所以只給 minor,
    # 真正的把關靠 --forbid(blocker)。
    rep.ran('stale_client_auto')
    orgs = extract_orgs(text)
    if not orgs:
        return
    rep.meta['出現的機構名'] = ', '.join(f'{k}×{v}' for k, v in
                                   sorted(orgs.items(), key=lambda x: -x[1]))
    main = own_name or max(orgs, key=orgs.get)
    allowed = set(DEFAULT_ALLOW) | {a.strip() for a in allow if a.strip()}
    for k, v in orgs.items():
        if k == main or k in main or main in k or k in allowed:
            continue
        rep.add('stale_client_auto', 'minor',
                f'文中出現非本案機構名「{k}」×{v}(本案判定為「{main}」)',
                fix=f'確認「{k}」是引用還是母版殘影;若是引用,加入 --allow 抑制')


def _check_placeholder(d, rep):
    rep.ran('placeholder')
    for p in all_paragraphs(d):
        m = PLACEHOLDER_RE.search(p.text)
        if m:
            rep.add('placeholder', 'major', f'殘留待填佔位字「{m.group()}」',
                    where=repr(p.text.strip()[:40]),
                    fix='補上實際內容,或在回報中明列為待補項目')


def _check_halfwidth_comma(d, rep):
    rep.ran('halfwidth_comma')
    for p in all_paragraphs(d):
        t = p.text
        for m in re.finditer(',', t):
            i = m.start()
            l = t[i - 1] if i else ''
            r = t[i + 1] if i + 1 < len(t) else ''
            if l.isdigit() and r.isdigit():
                continue          # 千分位,合法
            if l.isascii() and r.isascii():
                continue          # 純英文句,合法
            rep.add('halfwidth_comma', 'minor', '中文敘述中出現半形逗號',
                    where=repr(t[max(0, i - 20):i + 20]),
                    fix='改成全形「，」')


def _check_font_uniformity(d, rep):
    rep.ran('font_uniformity')
    for ti, t in enumerate(d.tables):
        combos = {}
        seen = []
        for row in t.rows:
            for c in row.cells:
                if any(p is c._tc for p in seen):
                    continue
                seen.append(c._tc)
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.text.strip():
                            combos[run_fmt(r)] = combos.get(run_fmt(r), 0) + 1
        if len(combos) > 1:
            desc = ' / '.join(f'{a or "繼承"}+{b or "繼承"}@{c or "繼承"}pt×{n}'
                              for (a, b, c), n in
                              sorted(combos.items(), key=lambda x: -x[1]))
            rep.add('font_uniformity', 'minor',
                    f'表格 #{ti} 內有 {len(combos)} 種字型/字級組合',
                    where=desc,
                    fix=f'把表格 #{ti} 所有 run 統一成佔多數的那組格式')


def _check_heading_style(d, rep):
    rep.ran('heading_style')
    pat = re.compile(r'^(?:[壹貳貮參叁肆伍陸陆柒捌玖拾]、|附件[一二三四五六七八九十]、)')
    heads = [p for p in d.paragraphs if pat.match(p.text.strip())]
    if not heads:
        return rep.skip('heading_style', '找不到「壹、」形式的章標題')
    styles = {}
    for p in heads:
        styles.setdefault(p.style.name, []).append(p.text.strip()[:18])
    rep.meta['章標題數'] = len(heads)
    if len(styles) > 1:
        rep.add('heading_style', 'major',
                f'章標題掛了 {len(styles)} 種段落樣式,更新目錄後會有章消失或跑錯階層',
                where='; '.join(f'{k}: {v}' for k, v in styles.items()),
                fix='把所有章標題統一掛同一個樣式(取佔多數者),文字與字型不要動')
    # 該樣式必須有 outlineLvl,否則不會進目錄
    for sname in styles:
        st = d.styles[sname]
        ppr = st.element.find(qn('w:pPr'))
        ol = ppr.find(qn('w:outlineLvl')) if ppr is not None else None
        if ol is None:
            rep.add('heading_outline', 'major',
                    f'章標題樣式「{sname}」沒有大綱階層(outlineLvl),不會進目錄',
                    fix=f'替「{sname}」設定 outlineLvl=0,或改掛有階層的標題樣式')
        else:
            rep.meta[f'樣式 {sname} 的大綱階層'] = ol.get(qn('w:val'))


def _check_toc(path, rep):
    rep.ran('toc_field')
    xml = zipfile.ZipFile(path).read('word/document.xml').decode('utf-8')
    m = re.search(r'TOC\s+\\o\s+"(\d)-(\d)"', xml)
    if not m:
        return rep.skip('toc_field', '文件沒有 TOC 網域')
    rep.meta['目錄範圍'] = f'\\o "{m.group(1)}-{m.group(2)}"'


def _check_measures(T, rep, min_chars):
    """逐張改善措施建議表:數字自洽 + 現況說明/改善方案字數。回傳各表摘要值。"""
    rep.ran('measure_math')
    rep.ran('section_min_chars')
    out = []
    for mi, t in enumerate(T['measures'], 1):
        name = t.cell(1, 1).text.strip()
        vals = {i: txt for i, txt in cells_of(t.rows[5])}
        use = parse_num(vals.get(0))          # 本項耗能量 kWh/年
        save = parse_num(vals.get(2))         # 節電量 kWh/年
        money = parse_num(vals.get(4))        # 節能效益 萬元/年
        rate = parse_num(vals.get(5))         # 節能率 %
        inv = parse_num(vals.get(9))          # 投資費用 萬元
        pay = parse_num(vals.get(10))         # 回收年限 年
        out.append(dict(name=name, save=save, money=money, inv=inv, pay=pay))
        tag = f'改善措施建議表({"一二三四五六七八九十"[mi-1]}) {name[:20]}'

        # 節電量欄不得有小數
        if save is not None and isinstance(save, float):
            rep.add('integer_kwh', 'minor', '節電量欄出現小數',
                    where=tag, fix='耗電量/節電量一律取整數並加千分位逗號')
        # 節能率 = 節電量 / 本項耗能量
        if use and save is not None and rate is not None and use > 0:
            calc = save / use * 100
            if not close(calc, rate, 0.02):
                rep.add('rate_mismatch', 'major',
                        f'節能率 {rate}% 與「節電量÷本項耗能量」推得的 {calc:.2f}% 不符',
                        where=tag, fix=f'把節能率改為 {calc:.2f}%,或修正節電量/本項耗能量')
        # 回收年限 = 投資 / 節費
        if inv and money and pay is not None and money > 0:
            calc = inv / money
            if not close(calc, pay, 0.02):
                rep.add('payback_mismatch', 'major',
                        f'回收年限 {pay} 年與「投資{inv}萬÷效益{money}萬」推得的 {calc:.2f} 年不符',
                        where=tag, fix=f'把回收年限改為 {calc:.2f} 年,或修正投資費用/節能效益')
        elif inv and money == 0 and pay is not None:
            rep.add('payback_undefined', 'minor',
                    f'節能效益為 0 卻填了回收年限「{vals.get(10)}」',
                    where=tag, fix='節能效益為 0 時回收年限請填「－」或「不適用」')

        # 現況說明 / 改善方案 字數
        body = t.cell(6, 0).text
        secs = re.split(r'\n(?=[一二三]、)', body)
        for label, key in (('現況說明', '一、'), ('改善方案', '二、')):
            sec = next((s for s in secs if s.startswith(key)), None)
            if sec is None:
                rep.add('section_missing', 'major', f'找不到「{key}{label}」段',
                        where=tag, fix=f'補寫「{key}{label}」段')
                continue
            n = len(''.join(sec.split('\n')[1:]).strip())
            if n < min_chars:
                rep.add('section_min_chars', 'major',
                        f'{label}只有 {n} 字,未達 {min_chars} 字下限',
                        where=tag,
                        fix=f'補實質內容(量測依據、規格、成因、作法)到 {min_chars} 字以上,不要灌水')
    return out


def _check_summary(T, rep, measures):
    """執行摘要表:各列 vs 改善措施表、小計 vs 各列加總。"""
    t = T['summary']
    if t is None:
        return rep.skip('summary_consistency', '找不到節能建議改善措施表')
    rep.ran('summary_consistency')
    rep.ran('summary_subtotal')

    # 找出資料列(項次為純數字)與小計列
    rows, subtotal = [], None
    for r in t.rows:
        cs = {i: txt for i, txt in cells_of(r)}
        first = cs.get(0, '')
        if first.isdigit():
            rows.append(cs)
        elif first.strip() in ('小計', '合計', '總計'):
            subtotal = cs

    # 逐列比對改善措施表
    for i, cs in enumerate(rows):
        if i >= len(measures):
            break
        m = measures[i]
        for col, key, label in ((2, 'save', '節電量'), (4, 'inv', '投資費用'),
                                (5, 'pay', '回收年限'), (3, 'money', '節能效益')):
            a, b = parse_num(cs.get(col)), m[key]
            if a is None and b is None:
                continue
            if not close(a, b, 0.005):
                rep.add('summary_consistency', 'blocker',
                        f'摘要表第{i+1}項的{label} {cs.get(col)!r} '
                        f'與改善措施建議表的 {b!r} 不一致',
                        where=f'{cs.get(1, "")[:24]}',
                        fix=f'兩處擇一為準並改成相同數值(改善措施表為計算來源,通常以它為準)')

    # 小計 = 各列加總
    if subtotal:
        for col, label, dec in ((2, '節電量', 0), (3, '節能效益', 2), (4, '投資費用', 0)):
            vals = [parse_num(cs.get(col)) for cs in rows]
            if any(v is None for v in vals):
                continue
            s = round(sum(vals), dec)
            got = parse_num(subtotal.get(col))
            if not close(got, s, 0.005):
                rep.add('summary_subtotal', 'blocker',
                        f'小計的{label} {got!r} 與各項加總 {s!r} 不符',
                        fix=f'把小計{label}改為 {s}')
        # 加權回收年限 = 總投資 / 總效益
        inv, money, pay = (parse_num(subtotal.get(4)), parse_num(subtotal.get(3)),
                           parse_num(subtotal.get(5)))
        if inv and money and pay is not None and money > 0:
            calc = inv / money
            if not close(calc, pay, 0.02):
                rep.add('summary_subtotal_payback', 'major',
                        f'小計回收年限 {pay} 與「總投資{inv}÷總效益{money}」推得的 {calc:.2f} 不符',
                        fix=f'把小計回收年限改為 {calc:.2f}')
